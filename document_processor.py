import re
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

CHUNK_SIZE    = 800
CHUNK_OVERLAP = 150


class DocumentProcessor:

    def extract_text(self, file_path: Path, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        extractors = {
            ".pdf":  self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt":  self._extract_txt,
            ".md":   self._extract_txt,
            ".csv":  self._extract_csv,
            ".xlsx": self._extract_xlsx,
        }
        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported file type: {ext}")
        try:
            text = extractor(file_path)
            return self._clean_text(text)
        except Exception as exc:
            raise RuntimeError(f"Could not extract text from '{filename}': {exc}")

    def split_into_chunks(self, text: str) -> List[str]:
        if not text.strip():
            return []
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        chunks: List[str] = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) + 2 <= CHUNK_SIZE:
                current = f"{current}\n\n{para}".strip()
            else:
                if current:
                    chunks.append(current)
                if len(para) > CHUNK_SIZE:
                    chunks.extend(self._split_long_text(para))
                    current = ""
                else:
                    current = para
        if current:
            chunks.append(current)
        overlapped = []
        for i, chunk in enumerate(chunks):
            if i > 0:
                chunk = chunks[i-1][-CHUNK_OVERLAP:] + "\n" + chunk
            overlapped.append(chunk)
        return [c for c in overlapped if len(c.strip()) > 50]

    def _extract_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
            if texts:
                return "\n\n".join(texts)
        except Exception:
            pass
        try:
            import PyPDF2
            texts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
            return "\n\n".join(texts)
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {e}")

    def _extract_docx(self, path: Path) -> str:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)

    def _extract_txt(self, path: Path) -> str:
        for enc in ("utf-8", "utf-16", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise RuntimeError("Cannot decode text file.")

    def _extract_csv(self, path: Path) -> str:
        import csv
        rows = []
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                rows.append(", ".join(row))
        return "\n".join(rows)

    def _extract_xlsx(self, path: Path) -> str:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r",   "\n", text)
        text = re.sub(r"\t",   " ",  text)
        text = re.sub(r" {3,}", "  ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        # Remove cid artifacts from PDFs
        text = re.sub(r"\(cid:\d+\)", "", text)
        text = re.sub(r"[^\x20-\x7E\u00A0-\uFFFF\n]", "", text)
        return text.strip()

    def _split_long_text(self, text: str) -> List[str]:
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks, current = [], ""
        for sent in sentences:
            if len(current) + len(sent) + 1 <= CHUNK_SIZE:
                current = f"{current} {sent}".strip()
            else:
                if current:
                    chunks.append(current)
                current = sent
        if current:
            chunks.append(current)
        return chunks
